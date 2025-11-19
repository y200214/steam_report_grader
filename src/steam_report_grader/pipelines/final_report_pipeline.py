# src/steam_report_grader/pipelines/final_report_pipeline.py
from __future__ import annotations

from pathlib import Path
import logging
from typing import List
from docx import Document

import pandas as pd

from ..utils.logging_utils import setup_logging
from ..config import AI_SUSPECT_THRESHOLD

logger = logging.getLogger(__name__)


def _build_final_results(
    absolute_scores_csv: Path,
    id_map_excel: Path,
) -> pd.DataFrame:
    """
    absolute_scores.csv + id_map.xlsx から final_results DataFrame を構築する。
    1 受験者 1 行。
    """
    absolute_scores_csv = Path(absolute_scores_csv)
    id_map_excel = Path(id_map_excel)

    scores_df = pd.read_csv(absolute_scores_csv)
    id_df = pd.read_excel(id_map_excel, sheet_name="id_map")

    if scores_df.empty:
        raise ValueError("absolute_scores_csv が空です")

    # ---- 設問スコアを横持ちにする（Q1, Q2 ... → Q1_score, Q2_score ...） ----
    # question 列には "Q1", "Q2" ... が入っている想定
    pivot = scores_df.pivot_table(
        index="student_id",
        columns="question",
        values="score",
        aggfunc="first",
    )

    # Q の番号順にソート
    def _q_key(q: str) -> int:
        try:
            return int(str(q).lstrip("Q"))
        except Exception:
            return 9999

    pivot = pivot.reindex(sorted(pivot.columns, key=_q_key), axis=1)

    # 列名を Q1 → Q1_score に揃える
    pivot = pivot.rename(columns={col: f"{col}_score" for col in pivot.columns})

    # 合計点
    pivot["total_score"] = pivot.sum(axis=1)

    final_df = pivot.reset_index()  # student_id が列に戻る

    # ---- id_map と結合して real_name / source_file を付与 ----
    final_df = final_df.merge(id_df, on="student_id", how="left")

    # 列順を整理
    score_cols: List[str] = [
        c for c in final_df.columns
        if c.endswith("_score") and c != "total_score"
    ]
    meta_cols: List[str] = [
        c for c in ["student_id", "real_name", "source_file"]
        if c in final_df.columns
    ]
    other_cols: List[str] = [
        c for c in final_df.columns
        if c not in meta_cols + ["total_score"] + score_cols
    ]

    ordered_cols = meta_cols + ["total_score"] + score_cols + other_cols
    final_df = final_df[ordered_cols]

    # ---- 順位計算（高得点が rank=1） ----
    final_df["rank"] = (
        final_df["total_score"]
        .rank(method="dense", ascending=False)
        .astype("Int64")
    )

    # 表示順も rank → student_id
    final_df = final_df.sort_values(["rank", "student_id"])

    return final_df

def _generate_feedback_markdown_for_student(
    scores_df: pd.DataFrame,
    final_row: pd.Series,
    questions_order: List[str],
) -> str:
    """
    ある受験者 1 人分の feedback_xxx.md の本文を組み立てる。
    """
    sid = final_row["student_id"]
    name = final_row.get("real_name", "") or "受験者"
    total = final_row.get("total_score", "")
    rank = final_row.get("rank", "")
    num_students = final_row.get("_num_students", None)

    # この受験者の全設問行を抜き出す
    s_rows = scores_df[scores_df["student_id"] == sid].copy()

    # 設問順を保証
    def _order_idx(q: str) -> int:
        try:
            return questions_order.index(q)
        except ValueError:
            return 9999

    s_rows["__q_order"] = s_rows["question"].apply(_order_idx)
    s_rows = s_rows.sort_values("__q_order")

    lines: List[str] = []

    # ヘッダー
    lines.append(f"# {name} さん (ID: {sid})")
    lines.append("")
    lines.append(f"- 総合点: {total}")
    if num_students is not None:
        lines.append(f"- 順位: {rank} / {num_students}")
    else:
        lines.append(f"- 順位: {rank}")

    # ← 相対順位を追記（あれば）
    relative_rank = final_row.get("relative_rank", None)
    if pd.notna(relative_rank):
        lines.append(f"- 相対順位: {int(relative_rank)}位")

    lines.append("")


    # 設問ごとの講評
    for _, row in s_rows.iterrows():
        q_label = row["question"]
        score = row["score"]
        # scoring_pipeline で brief カラムを入れているのでそれを優先
        brief = row.get("brief")
        if not isinstance(brief, str) or not brief.strip():
            brief = row.get("reason", "")

        lines.append(f"## {q_label} の講評（{score} 点）")
        lines.append("")
        if isinstance(brief, str) and brief.strip():
            lines.append(brief.strip())
        else:
            lines.append("※ この設問についての講評はありません。")
        lines.append("")

    # 将来的に総評などを追加するスペース
    # lines.append("## 総評")
    # lines.append("...")

    return "\n".join(lines)

def run_final_report(
    absolute_scores_csv: Path,
    id_map_excel: Path,
    ranking_csv_path: Path,
    feedback_dir: Path,
    log_path: Path,
) -> None:
    """
    - absolute_scores.csv + id_map.xlsx から final_results を構築
    - ranking.csv を出力
    - 各受験者ごとに feedback_{student_id}.md を生成

    ※ 現時点では絶対評価のみで順位付け（相対評価は未統合）
    """
    setup_logging(log_path)
    logger.info("Start final report pipeline")

    absolute_scores_csv = Path(absolute_scores_csv)
    id_map_excel = Path(id_map_excel)
    ranking_csv_path = Path(ranking_csv_path)
    feedback_dir = Path(feedback_dir)
    ranking_csv_path = Path(ranking_csv_path)

    ranking_csv_path.parent.mkdir(parents=True, exist_ok=True)

    # ルートだけ一応作っておく
    feedback_dir.mkdir(parents=True, exist_ok=True)

    # md / docx 用のサブフォルダ
    feedback_md_dir = feedback_dir / "md"
    feedback_docx_dir = feedback_dir / "docx"
    feedback_md_dir.mkdir(parents=True, exist_ok=True)
    feedback_docx_dir.mkdir(parents=True, exist_ok=True)


    # final_results を構築
    final_df = _build_final_results(
        absolute_scores_csv=absolute_scores_csv,
        id_map_excel=id_map_excel,
    )

    # ranking.csv として保存
    final_df.to_csv(ranking_csv_path, index=False, encoding="utf-8-sig")
    logger.info("Wrote ranking.csv to %s", ranking_csv_path)

    # feedback_xxx.md 生成のため、元の scores も読む
    scores_df = pd.read_csv(absolute_scores_csv)
    id_df = pd.read_excel(id_map_excel, sheet_name="id_map")
    # --- AI類似度 (ai_likeness.csv) を読み込む ---
    ai_likeness_path = Path("data/intermediate/features/ai_likeness.csv")
    ai_likeness_df: pd.DataFrame | None = None
    if ai_likeness_path.exists():
        try:
            tmp = pd.read_csv(ai_likeness_path)
            if not tmp.empty:
                ai_likeness_df = tmp
            logger.info(
                "Loaded AI likeness features from %s (rows=%d)",
                ai_likeness_path,
                0 if ai_likeness_df is None else len(ai_likeness_df),
            )
        except Exception as e:
            logger.warning("Failed to load ai_likeness.csv: %s", e)
    else:
        logger.warning(
            "AI likeness file %s not found. Skipping AI suspicion features.",
            ai_likeness_path,
        )

    # --- final_df に AI疑惑スコアをマージ ---
    if ai_likeness_df is not None and not ai_likeness_df.empty:
        agg = (
            ai_likeness_df
            .groupby("student_id", as_index=False)["ai_likeness_score"]
            .agg(
                ai_likeness_mean="mean",
                ai_likeness_max="max",
            )
        )

        # 少数第3位くらいで丸める
        agg["ai_likeness_mean"] = agg["ai_likeness_mean"].round(3)
        agg["ai_likeness_max"] = agg["ai_likeness_max"].round(3)

        # 閾値はひとまず 0.7（後で変えやすいようにここだけ）
        THRESHOLD = AI_SUSPECT_THRESHOLD
        agg["ai_suspect_flag"] = (agg["ai_likeness_max"] >= THRESHOLD).astype(int)

        # 型合わせしてマージ
        final_df["student_id"] = final_df["student_id"].astype(str)
        agg["student_id"] = agg["student_id"].astype(str)

        final_df = final_df.merge(agg, on="student_id", how="left")
        logger.info(
            "Merged AI likeness aggregates into final_df (cols added: %s)",
            ["ai_likeness_mean", "ai_likeness_max", "ai_suspect_flag"],
        )


    final_excel_path = ranking_csv_path.parent / "final_report.xlsx"
    _write_final_excel(
        final_df=final_df,
        scores_df=scores_df,
        id_df=id_df,
        ai_likeness_df=ai_likeness_df,
        excel_path=final_excel_path,
    )
    logger.info("Wrote final_report.xlsx to %s", final_excel_path)


    # 設問ラベルの順序（Q1, Q2, ...）を決める
    questions = sorted(
        scores_df["question"].unique().tolist(),
        key=lambda q: int(str(q).lstrip("Q")) if str(q).startswith("Q") else 9999,
    )

    # 設問ラベルの順序（Q1, Q2, ...）を決める
    questions = sorted(
        scores_df["question"].unique().tolist(),
        key=lambda q: int(str(q).lstrip("Q")) if str(q).startswith("Q") else 9999,
    )

    num_students = len(final_df)
    logger.info(
        "Generating feedback markdown/docx for %d students",
        num_students,
    )

    for _, row in final_df.iterrows():
        sid = row["student_id"]
        name = row.get("real_name", "") or "受験者"
        total = float(row.get("total_score", 0))
        rank = int(row.get("rank", 0))

        row_with_meta = row.copy()
        row_with_meta["_num_students"] = num_students

        md_text = _generate_feedback_markdown_for_student(
            scores_df=scores_df,
            final_row=row_with_meta,
            questions_order=questions,
        )

        # --- Markdown 出力（feedback/md/） ---
        md_path = feedback_md_dir / f"feedback_{sid}.md"
        md_path.write_text(md_text, encoding="utf-8")
        logger.info("Wrote feedback markdown for %s to %s", sid, md_path)

        # --- docx 出力（feedback/docx/） ---
        per_student_df = (
            scores_df[scores_df["student_id"] == sid]
            .sort_values("question")
        )
        mean = float(per_student_df["score"].mean())

        docx_path = feedback_docx_dir / f"feedback_{sid}.docx"
        _write_feedback_docx(
            student_id=str(sid),
            real_name=name if name != "受験者" else "",
            rank=rank,
            n_students=num_students,
            total=total,
            mean=mean,
            per_student_df=per_student_df,
            out_path=docx_path,
        )
        logger.info("Wrote feedback docx for %s to %s", sid, docx_path)


    logger.info("Final report pipeline completed.")

def _write_feedback_docx(
    student_id: str,
    real_name: str,
    rank: int,
    n_students: int,
    total: float,
    mean: float,
    per_student_df: pd.DataFrame,
    out_path: Path,
) -> None:
    """
    feedback_{student_id}.docx を書き出す。
    md と同じ情報を Word 形式で整形。
    """
    doc = Document()

    # タイトル
    title = f"{real_name} ({student_id})" if real_name else student_id
    doc.add_heading(title, level=1)

    # 基本情報
    p = doc.add_paragraph()
    p.add_run(f"順位: {rank} / {n_students}\n")
    p.add_run(f"合計得点: {total:.2f}\n")
    p.add_run(f"平均得点: {mean:.2f}")

    relative_rank = per_student_df["relative_rank"].iloc[0] if "relative_rank" in per_student_df.columns else None
    if relative_rank is not None:
        p.add_run(f"相対順位: {int(relative_rank)}\n")

    # 設問ごとの講評
    for _, srow in per_student_df.iterrows():
        q = srow["question"]
        score = srow["score"]

        brief = (
            srow.get("brief")
            or srow.get("brief_explanation")
            or srow.get("reason")
            or ""
        )
        detailed = (
            srow.get("detailed")
            or srow.get("detailed_explanation")
            or ""
        )



        doc.add_heading(f"{q} （{score} 点）", level=2)

        if isinstance(brief, str) and brief.strip():
            doc.add_heading("講評（簡易）", level=3)
            doc.add_paragraph(brief.strip())

        if isinstance(detailed, str) and detailed.strip():
            doc.add_heading("講評（詳細）", level=3)
            doc.add_paragraph(detailed.strip())

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)

def _write_final_excel(
    final_df: pd.DataFrame,
    scores_df: pd.DataFrame,
    id_df: pd.DataFrame,
    ai_likeness_df: pd.DataFrame | None,
    excel_path: Path,
) -> None:
    """
    final_results + scores を 1つの Excel にまとめて保存する。

    - Sheet 'ranking' : 1行 = 1受験者 (順位付き)
    - Sheet 'by_student_question' : 1行 = 1受験者×1設問
    """
    excel_path.parent.mkdir(parents=True, exist_ok=True)

    # 1. ranking シート用: final_df をそのまま
    ranking_df = final_df.copy()

    # 2. by_student_question シート用: scores_df + id_map を縦持ちで
    spq = scores_df.merge(id_df, on="student_id", how="left")
    
    # --- AI類似度を設問単位でマージ ---
    if ai_likeness_df is not None and not ai_likeness_df.empty:
        tmp = ai_likeness_df.copy()
        # 型合わせ
        if "student_id" in tmp.columns:
            tmp["student_id"] = tmp["student_id"].astype(str)
        if "question" in tmp.columns:
            tmp["question"] = tmp["question"].astype(str)

        spq["student_id"] = spq["student_id"].astype(str)
        spq["question"] = spq["question"].astype(str)

        cols_merge = [
            c for c in [
                "student_id",
                "question",
                "ai_likeness_score",
                "ai_likeness_comment",
            ]
            if c in tmp.columns
        ]
        if set(["student_id", "question"]).issubset(cols_merge):
            spq = spq.merge(
                tmp[cols_merge],
                on=["student_id", "question"],
                how="left",
            )

    # 必要な列だけに絞る
    cols = []
    for c in [
        "student_id",
        "real_name",
        "source_file",
        "question",
        "score",
        "brief",
        "detailed",
        "ai_likeness_score",
        "ai_likeness_comment",
    ]:
        if c in spq.columns:
            cols.append(c)
    spq = spq[cols]

    with pd.ExcelWriter(excel_path, engine="openpyxl") as writer:
        ranking_df.to_excel(writer, index=False, sheet_name="ranking")
        spq.to_excel(writer, index=False, sheet_name="by_student_question")

